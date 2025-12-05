from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import PyPDF2
import io
import os
import sys
from werkzeug.utils import secure_filename
import tempfile
import zipfile
try:
    from pdf2docx import Converter
    HAS_PDF2DOCX = True
except ImportError:
    HAS_PDF2DOCX = False

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

import google.generativeai as genai
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
CORS(app)  # Enable CORS for React frontend

@app.route('/')
def index():
    return jsonify({"message": "PDF API is running. Use /api/health to check status."})

UPLOAD_FOLDER = tempfile.gettempdir()
ALLOWED_EXTENSIONS = {'pdf'}
ALLOWED_OCR_EXTENSIONS = {'pdf', 'jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_ocr_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_OCR_EXTENSIONS

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({"status": "ok", "message": "PDF API is running"})

@app.route('/api/pdf/merge', methods=['POST'])
def merge_pdfs():
    """Merge multiple PDF files into one"""
    try:
        if 'files' not in request.files:
            return jsonify({"error": "No files provided"}), 400
        
        files = request.files.getlist('files')
        
        if len(files) < 2:
            return jsonify({"error": "At least 2 files required"}), 400
        
        print(f"Merging {len(files)} files...", flush=True)
        
        output = io.BytesIO()

        if HAS_PYMUPDF:
            print("Using PyMuPDF for merging", flush=True)
            merged_doc = fitz.open()
            for file in files:
                if file and allowed_file(file.filename):
                    print(f"Adding file: {file.filename}", flush=True)
                    file.seek(0)
                    file_bytes = file.read()
                    try:
                        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                            merged_doc.insert_pdf(doc)
                    except Exception as e:
                        print(f"Error adding file {file.filename}: {e}", flush=True)
                        return jsonify({"error": f"Failed to process {file.filename}: {str(e)}"}), 400
            
            merged_doc.save(output)
            merged_doc.close()
        else:
            print("Using PyPDF2 for merging", flush=True)
            # Create PDF merger
            merger = PyPDF2.PdfMerger()
            
            # Add all PDFs
            for file in files:
                if file and allowed_file(file.filename):
                    print(f"Adding file: {file.filename}", flush=True)
                    # Reset file pointer to ensure we read from start
                    file.seek(0)
                    merger.append(io.BytesIO(file.read()))
            
            merger.write(output)
            merger.close()
        
        output.seek(0)
        
        # Check output size
        size = output.getbuffer().nbytes
        print(f"Merge complete. Output size: {size} bytes", flush=True)
        
        if size == 0:
            return jsonify({"error": "Merged PDF is empty"}), 500
        
        return send_file(
            output,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='merged.pdf'
        )
    
    except Exception as e:
        print(f"Error in merge_pdfs: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/api/pdf/split', methods=['POST'])
def split_pdf():
    """Split PDF into separate files for each page range"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        page_range = request.form.get('pages', '')  # e.g., "1-5,8,10-12"
        
        if not file or not allowed_file(file.filename):
            return jsonify({"error": "Invalid file"}), 400
        
        if not page_range:
            return jsonify({"error": "Please specify pages to split (e.g., '1-5,8,10-12' or just '1-5')"}), 400
        
        # Read PDF
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
        total_pages = len(pdf_reader.pages)
        
        # Parse page range and create individual PDFs for each range
        ranges = []
        for part in page_range.split(','):
            part = part.strip()
            if '-' in part:
                try:
                    start, end = part.split('-')
                    start_page = int(start.strip()) - 1  # Convert to 0-indexed
                    end_page = int(end.strip())  # Inclusive
                    if start_page < 0 or end_page > total_pages or start_page >= end_page:
                        return jsonify({"error": f"Invalid page range: {part}. Valid pages: 1-{total_pages}"}), 400
                    ranges.append((start_page, end_page))
                except ValueError:
                    return jsonify({"error": f"Invalid page range format: {part}"}), 400
            else:
                try:
                    page_num = int(part) - 1  # Convert to 0-indexed
                    if page_num < 0 or page_num >= total_pages:
                        return jsonify({"error": f"Invalid page number: {part}. Valid pages: 1-{total_pages}"}), 400
                    ranges.append((page_num, page_num + 1))
                except ValueError:
                    return jsonify({"error": f"Invalid page number format: {part}"}), 400
        
        # If only one range, return a single PDF
        if len(ranges) == 1:
            start, end = ranges[0]
            pdf_writer = PyPDF2.PdfWriter()
            for page_num in range(start, end):
                pdf_writer.add_page(pdf_reader.pages[page_num])
            
            output = io.BytesIO()
            pdf_writer.write(output)
            output.seek(0)
            
            return send_file(
                output,
                mimetype='application/pdf',
                as_attachment=True,
                download_name='split.pdf'
            )
        
        # If multiple ranges, create ZIP with multiple PDFs
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for idx, (start, end) in enumerate(ranges, 1):
                pdf_writer = PyPDF2.PdfWriter()
                for page_num in range(start, end):
                    pdf_writer.add_page(pdf_reader.pages[page_num])
                
                # Create filename for this split
                if start == end - 1:
                    filename = f'split_page_{start + 1}.pdf'
                else:
                    filename = f'split_pages_{start + 1}-{end}.pdf'
                
                # Write to ZIP
                pdf_buffer = io.BytesIO()
                pdf_writer.write(pdf_buffer)
                pdf_buffer.seek(0)
                zip_file.writestr(filename, pdf_buffer.getvalue())
        
        zip_buffer.seek(0)
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name='split_pdfs.zip'
        )
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/pdf/unlock', methods=['POST'])
def unlock_pdf():
    """Remove password protection from PDF"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        password = request.form.get('password', '')
        
        if not file or not allowed_file(file.filename):
            return jsonify({"error": "Invalid file"}), 400
        
        # Read PDF
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
        
        # Check if encrypted
        if pdf_reader.is_encrypted:
            if password:
                pdf_reader.decrypt(password)
            else:
                return jsonify({"error": "PDF is password protected. Please provide password."}), 400
        
        # Create unlocked PDF
        pdf_writer = PyPDF2.PdfWriter()
        for page in pdf_reader.pages:
            pdf_writer.add_page(page)
        
        # Create output
        output = io.BytesIO()
        pdf_writer.write(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='unlocked.pdf'
        )
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/pdf/to-excel', methods=['POST'])
def pdf_to_excel():
    """Convert PDF to Excel by extracting tables in proper format"""
    temp_pdf = None
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        
        if not file or not allowed_file(file.filename):
            return jsonify({"error": "Invalid file"}), 400
        
        # Save temporarily
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        file.save(temp_pdf.name)
        temp_pdf.close()
        
        # Try using PyMuPDF if available
        if HAS_PYMUPDF:
            from openpyxl import Workbook
            from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
            
            # Open and read PDF
            doc = fitz.open(temp_pdf.name)
            
            # Create workbook
            wb = Workbook()
            ws = wb.active
            ws.title = "Extracted Data"
            
            current_row = 1
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # Extract tables from all pages
            tables_found = False
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Extract tables with proper structure
                try:
                    tables = page.find_tables()
                    if tables.tables:
                        tables_found = True
                        for table_idx, table in enumerate(tables.tables):
                            extracted_rows = table.extract()
                            
                            # Write header with page info
                            if current_row > 1:
                                current_row += 1  # Add space between tables
                            
                            # Write table rows with proper columns
                            for row_idx, row_data in enumerate(extracted_rows):
                                for col_idx, cell_value in enumerate(row_data):
                                    cell = ws.cell(row=current_row, column=col_idx + 1)
                                    # Clean up cell value
                                    cell_text = str(cell_value).strip() if cell_value else ""
                                    cell.value = cell_text
                                    cell.alignment = Alignment(wrap_text=True, vertical='top', horizontal='left')
                                    cell.border = thin_border
                                    
                                    # Format header row (usually first row)
                                    if row_idx == 0:
                                        cell.font = Font(bold=True, color="FFFFFF")
                                        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                                
                                current_row += 1
                except Exception as e:
                    print(f"Table extraction error on page {page_num + 1}: {str(e)}")
            
            # If no tables found, extract text as structured data
            if not tables_found:
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    text = page.get_text()
                    
                    if text.strip():
                        # Add page header
                        if current_row > 1:
                            current_row += 1
                        
                        cell = ws.cell(row=current_row, column=1)
                        cell.value = f"Page {page_num + 1}"
                        cell.font = Font(bold=True, size=11)
                        cell.fill = PatternFill(start_color="E8E8E8", end_color="E8E8E8", fill_type="solid")
                        current_row += 1
                        
                        # Parse lines into columns based on whitespace
                        lines = text.split('\n')
                        for line in lines:
                            if line.strip():
                                # Split by tabs or multiple spaces to identify columns
                                if '\t' in line:
                                    columns = line.split('\t')
                                else:
                                    # Split by 2+ spaces
                                    columns = [col.strip() for col in line.split('  ') if col.strip()]
                                
                                # If no clear columns, use single column
                                if not columns or len(columns) == 1:
                                    columns = [line.strip()]
                                
                                # Write columns
                                for col_idx, col_text in enumerate(columns[:15]):  # Max 15 columns
                                    cell = ws.cell(row=current_row, column=col_idx + 1)
                                    cell.value = col_text
                                    cell.alignment = Alignment(wrap_text=True, vertical='top')
                                    cell.border = thin_border
                                
                                current_row += 1
            
            # Auto-adjust column widths
            for col_idx in range(1, 16):
                col_letter = chr(64 + col_idx)
                ws.column_dimensions[col_letter].width = 30
            
            # Close PDF document
            doc.close()
            
            # Save to BytesIO
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            
            return send_file(
                output,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name='converted.xlsx'
            )
        else:
            return jsonify({"error": "PDF conversion requires PyMuPDF library"}), 501
    
    except Exception as e:
        print(f"PDF to Excel error: {str(e)}")
        return jsonify({"error": f"Failed to convert PDF to Excel: {str(e)}"}), 500
    
    finally:
        # Clean up temp file
        if temp_pdf and os.path.exists(temp_pdf.name):
            try:
                os.unlink(temp_pdf.name)
            except:
                pass

@app.route('/api/pdf/to-word', methods=['POST'])
def pdf_to_word():
    """Convert PDF to Word document"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        
        if not file or not allowed_file(file.filename):
            return jsonify({"error": "Invalid file"}), 400
        
        # Save temporarily
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        file.save(temp_pdf.name)
        temp_pdf.close()
        
        try:
            # Try using pdf2docx if available
            if HAS_PDF2DOCX:
                temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
                temp_docx.close()
                
                # Convert PDF to DOCX
                cv = Converter(temp_pdf.name)
                cv.convert(temp_docx.name)
                cv.close()
                
                # Read the converted file
                with open(temp_docx.name, 'rb') as f:
                    output = io.BytesIO(f.read())
                
                output.seek(0)
                os.unlink(temp_docx.name)
            
            # Fallback to PyMuPDF text extraction + python-docx
            elif HAS_PYMUPDF:
                try:
                    from docx import Document
                    from docx.shared import Pt, RGBColor
                    
                    doc = fitz.open(temp_pdf.name)
                    docx = Document()
                    
                    for page_num, page in enumerate(doc, 1):
                        # Add page number heading
                        heading = docx.add_heading(f'Page {page_num}', level=2)
                        heading.runs[0].font.size = Pt(12)
                        heading.runs[0].font.color.rgb = RGBColor(0, 0, 128)
                        
                        # Add text content
                        text = page.get_text()
                        for line in text.split('\n'):
                            if line.strip():
                                docx.add_paragraph(line)
                    
                    doc.close()
                    
                    output = io.BytesIO()
                    docx.save(output)
                    output.seek(0)
                except ImportError:
                    return jsonify({"error": "PDF to Word requires python-docx library"}), 501
            else:
                return jsonify({"error": "PDF conversion requires either pdf2docx or PyMuPDF library"}), 501
        finally:
            os.unlink(temp_pdf.name)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='converted.docx'
        )
    
    except Exception as e:
        print(f"PDF to Word error: {str(e)}")
        return jsonify({"error": f"Failed to convert PDF to Word: {str(e)}"}), 500

@app.route('/api/pdf/compress', methods=['POST'])
def compress_pdf():
    """Compress PDF file"""
    if not HAS_PYMUPDF:
        return jsonify({"error": "Server missing required library (PyMuPDF). Please install it to use this feature."}), 501
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        
        if not file or not allowed_file(file.filename):
            return jsonify({"error": "Invalid file"}), 400
        
        # Save temporarily
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        file.save(temp_pdf.name)
        temp_pdf.close()
        
        # Open with PyMuPDF
        doc = fitz.open(temp_pdf.name)
        
        # Compress
        output = io.BytesIO()
        doc.save(output, garbage=4, deflate=True, clean=True)
        doc.close()
        
        output.seek(0)
        os.unlink(temp_pdf.name)
        
        return send_file(
            output,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='compressed.pdf'
        )
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/pdf/extract-text', methods=['POST'])
def extract_text():
    """Extract text from PDF"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        
        if not file or not allowed_file(file.filename):
            return jsonify({"error": "Invalid file"}), 400
        
        # Read PDF
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
        
        # Extract text from all pages
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n\n"
        
        return jsonify({
            "text": text,
            "pages": len(pdf_reader.pages)
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/pdf/info', methods=['POST'])
def pdf_info():
    """Get PDF information"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        
        if not file or not allowed_file(file.filename):
            return jsonify({"error": "Invalid file"}), 400
        
        # Read PDF
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
        
        info = {
            "pages": len(pdf_reader.pages),
            "encrypted": pdf_reader.is_encrypted,
            "metadata": {}
        }
        
        if pdf_reader.metadata:
            info["metadata"] = {
                "title": pdf_reader.metadata.get('/Title', 'N/A'),
                "author": pdf_reader.metadata.get('/Author', 'N/A'),
                "subject": pdf_reader.metadata.get('/Subject', 'N/A'),
                "creator": pdf_reader.metadata.get('/Creator', 'N/A'),
            }
        
        return jsonify(info)
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/pdf/add-watermark', methods=['POST'])
def add_watermark():
    """Add text watermark to PDF with 45 degree rotation"""
    if not HAS_PYMUPDF:
        return jsonify({"error": "Server missing required library (PyMuPDF). Please install it to use this feature."}), 501
    temp_pdf = None
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        watermark_text = request.form.get('text', 'CONFIDENTIAL')
        
        if not file or not allowed_file(file.filename):
            return jsonify({"error": "Invalid file"}), 400
        
        # Save temporarily
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        file.save(temp_pdf.name)
        temp_pdf.close()
        
        # Open with PyMuPDF
        doc = fitz.open(temp_pdf.name)
        
        # Add watermark to each page
        for page in doc:
            # Get page dimensions
            rect = page.rect
            center_x = rect.width / 2
            center_y = rect.height / 2
            
            # Create watermark text rotated 45 degrees
            # Draw text in upper-left to lower-right diagonal pattern
            import math
            
            # Calculate multiple diagonal lines to cover the page
            angle = math.radians(45)
            
            # Text properties
            font_size = 48
            text_color = (0.7, 0.7, 0.7)  # Light gray
            
            # Draw multiple watermark texts diagonally across page
            step = 200  # Distance between watermark repetitions
            
            # Start from upper left, go to lower right
            for x_offset in range(-int(rect.width), int(rect.width) + 1, step):
                for y_offset in range(-int(rect.height), int(rect.height) + 1, step):
                    # Position on diagonal
                    x = center_x + x_offset
                    y = center_y + y_offset
                    
                    if 0 <= x <= rect.width and 0 <= y <= rect.height:
                        # Use shape with text drawing
                        shape = page.new_shape()
                        shape.insert_text(
                            fitz.Point(x - 80, y - 15),
                            watermark_text,
                            fontsize=font_size,
                            color=text_color,
                            fontname="helv"
                        )
                        shape.commit(overlay=False)
        
        # Save to output
        output = io.BytesIO()
        doc.save(output)
        doc.close()
        
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='watermarked.pdf'
        )
    
    except Exception as e:
        print(f"Watermark error: {str(e)}")
        return jsonify({"error": f"Failed to add watermark: {str(e)}"}), 500
    
    finally:
        # Clean up temp file
        if temp_pdf and os.path.exists(temp_pdf.name):
            try:
                os.unlink(temp_pdf.name)
            except:
                pass

@app.route('/api/pdf/rotate', methods=['POST'])
def rotate_pdf():
    """Rotate PDF pages"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        rotation = int(request.form.get('rotation', 90))  # 90, 180, 270
        
        if not file or not allowed_file(file.filename):
            return jsonify({"error": "Invalid file"}), 400
        
        if rotation not in [90, 180, 270, -90]:
            return jsonify({"error": "Rotation must be 90, 180, or 270 degrees"}), 400
        
        # Read PDF
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
        pdf_writer = PyPDF2.PdfWriter()
        
        # Rotate all pages
        for page in pdf_reader.pages:
            page.rotate(rotation)
            pdf_writer.add_page(page)
        
        # Create output
        output = io.BytesIO()
        pdf_writer.write(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='rotated.pdf'
        )
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/pdf/extract-images', methods=['POST'])
def extract_images():
    """Extract all images from PDF"""
    if not HAS_PYMUPDF:
        return jsonify({"error": "Server missing required library (PyMuPDF). Please install it to use this feature."}), 501
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        
        if not file or not allowed_file(file.filename):
            return jsonify({"error": "Invalid file"}), 400
        
        # Save temporarily
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        file.save(temp_pdf.name)
        temp_pdf.close()
        
        # Open with PyMuPDF
        doc = fitz.open(temp_pdf.name)
        
        images = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                images.append({
                    "page": page_num + 1,
                    "index": img_index + 1,
                    "width": base_image["width"],
                    "height": base_image["height"],
                    "format": base_image["ext"]
                })
        
        doc.close()
        os.unlink(temp_pdf.name)
        
        return jsonify({
            "images": images,
            "total": len(images)
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/ocr/gemini', methods=['POST'])
def ocr_gemini():
    """Extract text using Gemini Vision API"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        # Try to get key from request first, then env
        api_key = request.form.get('api_key') or os.getenv('GEMINI_API_KEY')
        
        if not api_key:
            return jsonify({"error": "Gemini API Key is required. Please add GEMINI_API_KEY to .env or provide it in the request."}), 400
            
        if not file or file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        if not allowed_ocr_file(file.filename):
            return jsonify({"error": f"Invalid file format. Supported formats: {', '.join(ALLOWED_OCR_EXTENSIONS)}"}), 400
            
        # Configure Gemini
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        # Save temporarily
        temp_path = os.path.join(tempfile.gettempdir(), secure_filename(file.filename))
        file.save(temp_path)
        
        try:
            # Verify file was saved correctly
            if not os.path.exists(temp_path):
                return jsonify({"error": "Failed to save uploaded file"}), 400
            
            # Check file size
            file_size = os.path.getsize(temp_path)
            if file_size == 0:
                return jsonify({"error": "Uploaded file is empty"}), 400
            
            # Upload to Gemini
            uploaded_file = genai.upload_file(temp_path)
            
            # Generate content with mime type awareness
            prompt = "Extract all text from this document. Preserve the original formatting as much as possible. If there are tables, represent them with markdown tables. Do not include any introductory or concluding remarks, just the extracted text."
            
            response = model.generate_content([
                prompt,
                uploaded_file
            ])
            
            if not response.text:
                return jsonify({"error": "No text could be extracted from the document. The document may be empty or unreadable."}), 400
            
            # Cleanup remote file
            try:
                uploaded_file.delete()
            except Exception as cleanup_error:
                print(f"Warning: Failed to cleanup remote file: {cleanup_error}")
                pass
                
            return jsonify({
                "text": response.text,
            })
            
        except Exception as e:
            error_message = str(e)
            print(f"OCR Error: {error_message}")
            return jsonify({"error": f"OCR processing failed: {error_message}"}), 500
        finally:
            if os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except Exception as cleanup_error:
                    print(f"Warning: Failed to cleanup temp file: {cleanup_error}")
                
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    print("🚀 PDF API Server starting...")
    print("📍 Running on http://localhost:5000")
    app.run(debug=True, port=5000)
